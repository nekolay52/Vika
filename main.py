from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document
from docx.shared import Pt
from bs4 import BeautifulSoup
from docx.shared import Mm
import requests
import re


def dngffdfkg(doc):
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    return doc

while True:
    url = input("введите ссылку на статью Википедии: ")

    if "https://ru.wikipedia.org" not in url:
        print("ну мы тебя же просили типа ввести статью на Википедию что ты делаешь")
    else:
        break

headers = {
    "User-Agent":"Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/85.0.4183.102 YaBrowser/20.9.3.136 Yowser/2.5 Safari/537.36"
}

response = requests.get(url, headers=headers)
response.raise_for_status()

soup = BeautifulSoup(response.text, features="html.parser")
some_headline = soup.find('h1', class_="firstHeading mw-first-heading")
ocnovnoy = soup.find('div', class_="mw-content-ltr mw-parser-output")
tosh_divan = ocnovnoy.find_all(['p', "h1", "h2", "h3", "h4", "h5", "h6", "img"])
tosh_divan = [some_headline] + tosh_divan

#удоление пустЮ разд-ов

for i in range(len(tosh_divan) - 1, -1, -1):
    if "<p" in str(tosh_divan[i]):
        w = i
        break

qyqy = tosh_divan[:w + 1]

doc = Document()

doc = dngffdfkg(doc)


for teg in qyqy:
    if "<p>" in str(teg):
        w = re.sub(r'\[.*?\]', '', teg.text)
        s = doc.add_paragraph(w)
        s.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if "<h1" in str(teg):
        head = doc.add_heading(teg.text, level=1)
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if "<h2" in str(teg):
        head = doc.add_heading(teg.text, level=2)
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if "<h3" in str(teg):
        head = doc.add_heading(teg.text, level=3)
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if "<h4" in str(teg):
        head = doc.add_heading(teg.text, level=4)
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if "<h5" in str(teg):
        head = doc.add_heading(teg.text, level=5)
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if "<h6" in str(teg):
        head = doc.add_heading(teg.text, level=6)
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if "<img" in str(teg):
        response = requests.get("https:" + teg['src'], headers=headers)
        response.raise_for_status()
        with open('filename.png', "wb") as f:
            f.write(response.content)
        imga = doc.add_picture('filename.png')
        imga2 = doc.paragraphs[-1]
        imga2.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.save('test.docx')


